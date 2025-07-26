"""
Document manager using python-docx for file operations.

This module provides document reading and analysis capabilities using python-docx.
"""

import os
import logging
from typing import Dict, Any, List, Optional
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Inches

from ..utils.errors import DocumentError, OperationError
from ..utils.validators import validate_file_path


logger = logging.getLogger(__name__)


class DocumentManager:
    """Handles document operations using python-docx.
    
    This class provides functionality to read Word documents, extract content,
    analyze document structure, and extract metadata.
    """
    
    def __init__(self, config):
        """Initialize document manager.
        
        Args:
            config: Configuration object
        """
        self.config = config
        logger.info("DocumentManager initialized")
    
    def read_document(self, path: str) -> Dict[str, Any]:
        """Read document content and structure.
        
        Args:
            path: Document file path
            
        Returns:
            Dictionary containing document content, structure, and metadata
            
        Raises:
            DocumentError: If document cannot be read or is invalid
        """
        try:
            # Validate file path
            if not validate_file_path(path):
                raise DocumentError(f"Invalid file path: {path}")
            
            # Check if file exists
            if not os.path.exists(path):
                raise DocumentError(f"Document not found: {path}")
            
            # Check file extension
            file_path = Path(path)
            if file_path.suffix.lower() not in ['.docx', '.doc']:
                raise DocumentError(f"Unsupported file format: {file_path.suffix}")
            
            # Load document
            try:
                doc = Document(path)
            except Exception as e:
                raise DocumentError(f"Failed to load document: {str(e)}")
            
            # Extract content and structure
            result = {
                "path": path,
                "text": self._extract_text_from_doc(doc),
                "structure": self._analyze_document_structure(doc),
                "tables": self._extract_tables(doc),
                "paragraphs": self._extract_paragraphs(doc),
                "metadata": self._extract_basic_metadata(path, doc)
            }
            
            logger.info(f"Successfully read document: {path}")
            return result
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error reading document {path}: {e}")
            raise DocumentError(f"Unexpected error reading document: {str(e)}")
    
    def extract_text(self, path: str) -> str:
        """Extract plain text from document.
        
        Args:
            path: Document file path
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentError: If document cannot be read
        """
        try:
            # Validate file path
            if not validate_file_path(path):
                raise DocumentError(f"Invalid file path: {path}")
            
            # Check if file exists
            if not os.path.exists(path):
                raise DocumentError(f"Document not found: {path}")
            
            # Load document and extract text
            try:
                doc = Document(path)
                text = self._extract_text_from_doc(doc)
                
                logger.info(f"Extracted {len(text)} characters from document: {path}")
                return text
                
            except Exception as e:
                raise DocumentError(f"Failed to extract text from document: {str(e)}")
                
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error extracting text from {path}: {e}")
            raise DocumentError(f"Unexpected error extracting text: {str(e)}")
    
    def get_document_structure(self, path: str) -> Dict[str, Any]:
        """Analyze document structure (headings, paragraphs).
        
        Args:
            path: Document file path
            
        Returns:
            Dictionary containing document structure information
            
        Raises:
            DocumentError: If document cannot be analyzed
        """
        try:
            # Validate file path
            if not validate_file_path(path):
                raise DocumentError(f"Invalid file path: {path}")
            
            # Check if file exists
            if not os.path.exists(path):
                raise DocumentError(f"Document not found: {path}")
            
            # Load document and analyze structure
            try:
                doc = Document(path)
                structure = self._analyze_document_structure(doc)
                
                logger.info(f"Analyzed structure of document: {path}")
                return structure
                
            except Exception as e:
                raise DocumentError(f"Failed to analyze document structure: {str(e)}")
                
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error analyzing structure of {path}: {e}")
            raise DocumentError(f"Unexpected error analyzing structure: {str(e)}")
    
    def _extract_text_from_doc(self, doc: DocxDocument) -> str:
        """Extract all text from a document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Extracted text content
        """
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return '\n'.join(text_parts)
    
    def _analyze_document_structure(self, doc: DocxDocument) -> Dict[str, Any]:
        """Analyze document structure including headings and paragraphs.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Dictionary containing structure information
        """
        structure = {
            "headings": [],
            "paragraph_count": 0,
            "table_count": len(doc.tables),
            "sections": []
        }
        
        current_section = {
            "level": 0,
            "title": "Document Start",
            "paragraphs": []
        }
        
        for i, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph is a heading
            if paragraph.style.name.startswith('Heading'):
                # Extract heading level
                try:
                    level = int(paragraph.style.name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                
                heading_info = {
                    "level": level,
                    "text": paragraph.text,
                    "paragraph_index": i
                }
                structure["headings"].append(heading_info)
                
                # Start new section
                if current_section["paragraphs"]:
                    structure["sections"].append(current_section)
                
                current_section = {
                    "level": level,
                    "title": paragraph.text,
                    "paragraphs": []
                }
            else:
                # Regular paragraph
                if paragraph.text.strip():
                    current_section["paragraphs"].append({
                        "index": i,
                        "text": paragraph.text,
                        "style": paragraph.style.name
                    })
                    structure["paragraph_count"] += 1
        
        # Add final section
        if current_section["paragraphs"]:
            structure["sections"].append(current_section)
        
        return structure
    
    def _extract_paragraphs(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """Extract detailed paragraph information.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of paragraph information dictionaries
        """
        paragraphs = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            para_info = {
                "index": i,
                "text": paragraph.text,
                "style": paragraph.style.name,
                "alignment": str(paragraph.alignment) if paragraph.alignment else None,
                "runs": []
            }
            
            # Extract run information (formatting details)
            for run in paragraph.runs:
                run_info = {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name if run.font.name else None,
                    "font_size": run.font.size.pt if run.font.size else None
                }
                para_info["runs"].append(run_info)
            
            paragraphs.append(para_info)
        
        return paragraphs
    
    def _extract_tables(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """Extract table information from document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of table information dictionaries
        """
        tables = []
        
        for i, table in enumerate(doc.tables):
            table_info = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
                "data": []
            }
            
            # Extract table data
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell_idx, cell in enumerate(row.cells):
                    cell_data = {
                        "text": cell.text,
                        "row": row_idx,
                        "column": cell_idx
                    }
                    row_data.append(cell_data)
                table_info["data"].append(row_data)
            
            tables.append(table_info)
        
        return tables
    
    def _extract_basic_metadata(self, path: str, doc: DocxDocument) -> Dict[str, Any]:
        """Extract basic metadata from document.
        
        Args:
            path: Document file path
            doc: python-docx Document object
            
        Returns:
            Dictionary containing basic metadata
        """
        file_path = Path(path)
        file_stat = file_path.stat()
        
        metadata = {
            "filename": file_path.name,
            "file_size": file_stat.st_size,
            "created": file_stat.st_ctime,
            "modified": file_stat.st_mtime,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
        
        # Try to extract document properties
        try:
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["document_created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["document_modified"] = core_props.modified.isoformat()
        except Exception as e:
            logger.warning(f"Could not extract document properties: {e}")
        
        return metadata
    
    def get_document_info(self, path: str) -> Dict[str, Any]:
        """Get comprehensive document metadata and statistics.
        
        Args:
            path: Document file path
            
        Returns:
            Dictionary containing metadata and statistics
            
        Raises:
            DocumentError: If document cannot be analyzed
        """
        try:
            # Validate file path
            if not validate_file_path(path):
                raise DocumentError(f"Invalid file path: {path}")
            
            # Check if file exists
            if not os.path.exists(path):
                raise DocumentError(f"Document not found: {path}")
            
            # Load document
            try:
                doc = Document(path)
            except Exception as e:
                raise DocumentError(f"Failed to load document: {str(e)}")
            
            # Get basic metadata
            metadata = self._extract_basic_metadata(path, doc)
            
            # Add statistics
            statistics = self._calculate_document_statistics(doc)
            metadata.update(statistics)
            
            # Add document properties
            properties = self._extract_document_properties(doc)
            metadata.update(properties)
            
            logger.info(f"Extracted metadata and statistics for document: {path}")
            return metadata
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error getting document info for {path}: {e}")
            raise DocumentError(f"Unexpected error getting document info: {str(e)}")
    
    def get_document_statistics(self, path: str) -> Dict[str, Any]:
        """Get document statistics (word count, page count, etc.).
        
        Args:
            path: Document file path
            
        Returns:
            Dictionary containing document statistics
            
        Raises:
            DocumentError: If document cannot be analyzed
        """
        try:
            # Validate file path
            if not validate_file_path(path):
                raise DocumentError(f"Invalid file path: {path}")
            
            # Check if file exists
            if not os.path.exists(path):
                raise DocumentError(f"Document not found: {path}")
            
            # Load document
            try:
                doc = Document(path)
            except Exception as e:
                raise DocumentError(f"Failed to load document: {str(e)}")
            
            # Calculate statistics
            statistics = self._calculate_document_statistics(doc)
            
            logger.info(f"Calculated statistics for document: {path}")
            return statistics
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error calculating statistics for {path}: {e}")
            raise DocumentError(f"Unexpected error calculating statistics: {str(e)}")
    
    def extract_comments(self, path: str) -> List[Dict[str, Any]]:
        """Extract comments from document.
        
        Args:
            path: Document file path
            
        Returns:
            List of comment information dictionaries
            
        Raises:
            DocumentError: If document cannot be analyzed
        """
        try:
            # Validate file path
            if not validate_file_path(path):
                raise DocumentError(f"Invalid file path: {path}")
            
            # Check if file exists
            if not os.path.exists(path):
                raise DocumentError(f"Document not found: {path}")
            
            # Load document
            try:
                doc = Document(path)
            except Exception as e:
                raise DocumentError(f"Failed to load document: {str(e)}")
            
            # Extract comments (python-docx has limited comment support)
            comments = self._extract_comments_from_doc(doc)
            
            logger.info(f"Extracted {len(comments)} comments from document: {path}")
            return comments
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error extracting comments from {path}: {e}")
            raise DocumentError(f"Unexpected error extracting comments: {str(e)}")
    
    def _calculate_document_statistics(self, doc: DocxDocument) -> Dict[str, Any]:
        """Calculate comprehensive document statistics.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Dictionary containing statistics
        """
        # Extract all text for word counting
        all_text = self._extract_text_from_doc(doc)
        
        # Calculate word count
        words = all_text.split()
        word_count = len(words)
        
        # Calculate character count
        char_count = len(all_text)
        char_count_no_spaces = len(all_text.replace(' ', ''))
        
        # Count paragraphs (non-empty)
        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
        
        # Count sentences (rough estimate)
        sentence_count = len([s for s in all_text.split('.') if s.strip()])
        
        # Count tables
        table_count = len(doc.tables)
        
        # Count images/shapes (limited support in python-docx)
        image_count = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if hasattr(run, '_element') and run._element.xpath('.//a:blip'):
                    image_count += 1
        
        # Count headings by style
        heading_counts = {}
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name
                heading_counts[level] = heading_counts.get(level, 0) + 1
        
        # Estimate reading time (average 200 words per minute)
        reading_time_minutes = max(1, word_count // 200)
        
        statistics = {
            "word_count": word_count,
            "character_count": char_count,
            "character_count_no_spaces": char_count_no_spaces,
            "paragraph_count": paragraph_count,
            "sentence_count": sentence_count,
            "table_count": table_count,
            "image_count": image_count,
            "heading_counts": heading_counts,
            "total_headings": sum(heading_counts.values()),
            "reading_time_minutes": reading_time_minutes
        }
        
        return statistics
    
    def _extract_document_properties(self, doc: DocxDocument) -> Dict[str, Any]:
        """Extract comprehensive document properties.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Dictionary containing document properties
        """
        properties = {}
        
        try:
            core_props = doc.core_properties
            
            # Core properties
            if core_props.title:
                properties["title"] = core_props.title
            if core_props.author:
                properties["author"] = core_props.author
            if core_props.subject:
                properties["subject"] = core_props.subject
            if core_props.keywords:
                properties["keywords"] = core_props.keywords
            if core_props.comments:
                properties["comments"] = core_props.comments
            if core_props.category:
                properties["category"] = core_props.category
            if core_props.language:
                properties["language"] = core_props.language
            if core_props.created:
                properties["created"] = core_props.created.isoformat()
            if core_props.modified:
                properties["modified"] = core_props.modified.isoformat()
            if core_props.last_modified_by:
                properties["last_modified_by"] = core_props.last_modified_by
            if core_props.revision:
                properties["revision"] = core_props.revision
            if core_props.version:
                properties["version"] = core_props.version
                
        except Exception as e:
            logger.warning(f"Could not extract document properties: {e}")
        
        return properties
    
    def _extract_comments_from_doc(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """Extract comments from document (limited support in python-docx).
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of comment dictionaries
        """
        comments = []
        
        # Note: python-docx has limited support for comments
        # This is a basic implementation that may not capture all comments
        try:
            # Try to access comments through the document's XML
            if hasattr(doc, '_element'):
                # Look for comment references in the document
                comment_refs = doc._element.xpath('.//w:commentReference', 
                                                namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                
                for i, ref in enumerate(comment_refs):
                    comment_id = ref.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    comment_info = {
                        "id": comment_id or str(i),
                        "text": f"Comment {i + 1}",  # Placeholder - actual text extraction is complex
                        "author": "Unknown",  # Would need to parse comments.xml
                        "date": None,  # Would need to parse comments.xml
                        "position": i
                    }
                    comments.append(comment_info)
                    
        except Exception as e:
            logger.warning(f"Could not extract comments: {e}")
        
        return comments