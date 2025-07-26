#!/usr/bin/env python3
"""
Mac-compatible Word MCP Server using python-docx.
This version works on Mac, Linux, and Windows without requiring Microsoft Word.
"""

import asyncio
import json
import sys
import os
import uuid
from pathlib import Path
from typing import Any, Dict, List
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

class MacWordMCPServer:
    def __init__(self):
        self.tools = self._create_word_tools()
        self.documents = {}  # Store open documents
        self.output_dir = Path.home() / "Documents" / "MCP_Word_Documents"
        self.output_dir.mkdir(exist_ok=True)
        
    def _create_word_tools(self):
        """Create the list of Word tools that work on Mac."""
        return [
            {
                "name": "create_document",
                "description": "Create a new Word document (works on Mac using python-docx)",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "title": {
                            "type": "string",
                            "description": "Optional title for the document"
                        }
                    }
                }
            },
            {
                "name": "insert_text",
                "description": "Insert text into a Word document",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "doc_id": {"type": "string", "description": "Document ID"},
                        "text": {"type": "string", "description": "Text to insert"}
                    },
                    "required": ["doc_id", "text"]
                }
            },
            {
                "name": "format_text",
                "description": "Apply formatting to the last added text",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "doc_id": {"type": "string", "description": "Document ID"},
                        "bold": {"type": "boolean", "description": "Apply bold formatting"},
                        "italic": {"type": "boolean", "description": "Apply italic formatting"},
                        "font_size": {"type": "integer", "description": "Font size in points"}
                    },
                    "required": ["doc_id"]
                }
            },
            {
                "name": "add_heading",
                "description": "Add a heading to the document",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "doc_id": {"type": "string", "description": "Document ID"},
                        "text": {"type": "string", "description": "Heading text"},
                        "level": {"type": "integer", "description": "Heading level (1-3)", "minimum": 1, "maximum": 3}
                    },
                    "required": ["doc_id", "text"]
                }
            },
            {
                "name": "save_document",
                "description": "Save a Word document to ~/Documents/MCP_Word_Documents/",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "doc_id": {"type": "string", "description": "Document ID"},
                        "filename": {"type": "string", "description": "Filename (without path)"}
                    },
                    "required": ["doc_id", "filename"]
                }
            },
            {
                "name": "read_document",
                "description": "Read content from a Word document",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Path to document"}
                    },
                    "required": ["path"]
                }
            },
            {
                "name": "create_table",
                "description": "Create a table in a Word document",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "doc_id": {"type": "string", "description": "Document ID"},
                        "rows": {"type": "integer", "description": "Number of rows", "minimum": 1, "maximum": 20},
                        "cols": {"type": "integer", "description": "Number of columns", "minimum": 1, "maximum": 10}
                    },
                    "required": ["doc_id", "rows", "cols"]
                }
            },
            {
                "name": "create_list",
                "description": "Create a bulleted or numbered list",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "doc_id": {"type": "string", "description": "Document ID"},
                        "items": {"type": "array", "items": {"type": "string"}, "description": "List items"},
                        "list_type": {"type": "string", "enum": ["bulleted", "numbered"], "default": "bulleted"}
                    },
                    "required": ["doc_id", "items"]
                }
            },
            {
                "name": "list_documents",
                "description": "List all documents created by this MCP server",
                "inputSchema": {
                    "type": "object",
                    "properties": {}
                }
            }
        ]
    
    def _create_document(self, title: str = None) -> str:
        """Create a new document and return its ID."""
        doc_id = str(uuid.uuid4())
        doc = Document()
        
        if title:
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.documents[doc_id] = doc
        return doc_id
    
    def _insert_text(self, doc_id: str, text: str) -> str:
        """Insert text into a document."""
        if doc_id not in self.documents:
            return f"❌ Document {doc_id} not found"
        
        doc = self.documents[doc_id]
        paragraph = doc.add_paragraph(text)
        return f"✅ Added text to document: '{text[:50]}...'" if len(text) > 50 else f"✅ Added text: '{text}'"
    
    def _format_last_paragraph(self, doc_id: str, bold: bool = None, italic: bool = None, font_size: int = None) -> str:
        """Apply formatting to the last paragraph."""
        if doc_id not in self.documents:
            return f"❌ Document {doc_id} not found"
        
        doc = self.documents[doc_id]
        if not doc.paragraphs:
            return "❌ No paragraphs to format"
        
        last_paragraph = doc.paragraphs[-1]
        
        for run in last_paragraph.runs:
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if font_size is not None:
                run.font.size = Pt(font_size)
        
        formatting = []
        if bold: formatting.append("bold")
        if italic: formatting.append("italic")
        if font_size: formatting.append(f"{font_size}pt")
        
        return f"✅ Applied formatting: {', '.join(formatting) if formatting else 'none'}"
    
    def _add_heading(self, doc_id: str, text: str, level: int = 1) -> str:
        """Add a heading to the document."""
        if doc_id not in self.documents:
            return f"❌ Document {doc_id} not found"
        
        doc = self.documents[doc_id]
        doc.add_heading(text, level)
        return f"✅ Added heading (level {level}): '{text}'"
    
    def _save_document(self, doc_id: str, filename: str) -> str:
        """Save a document to file."""
        if doc_id not in self.documents:
            return f"❌ Document {doc_id} not found"
        
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        filepath = self.output_dir / filename
        doc = self.documents[doc_id]
        
        try:
            doc.save(str(filepath))
            return f"✅ Document saved to: {filepath}"
        except Exception as e:
            return f"❌ Failed to save document: {str(e)}"
    
    def _read_document(self, path: str) -> str:
        """Read content from a document."""
        try:
            doc = Document(path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            if not content:
                return "📄 Document is empty or contains no readable text"
            
            return f"📄 Document content ({len(content)} paragraphs):\n\n" + "\n\n".join(content)
        
        except Exception as e:
            return f"❌ Failed to read document: {str(e)}"
    
    def _create_table(self, doc_id: str, rows: int, cols: int) -> str:
        """Create a table in the document."""
        if doc_id not in self.documents:
            return f"❌ Document {doc_id} not found"
        
        doc = self.documents[doc_id]
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # Add sample headers
        if rows > 0 and cols > 0:
            for i in range(cols):
                table.cell(0, i).text = f"Header {i+1}"
        
        return f"✅ Created {rows}x{cols} table with sample headers"
    
    def _create_list(self, doc_id: str, items: List[str], list_type: str = "bulleted") -> str:
        """Create a list in the document."""
        if doc_id not in self.documents:
            return f"❌ Document {doc_id} not found"
        
        doc = self.documents[doc_id]
        
        for i, item in enumerate(items):
            if list_type == "numbered":
                paragraph = doc.add_paragraph(f"{i+1}. {item}")
            else:
                paragraph = doc.add_paragraph(item, style='List Bullet')
        
        return f"✅ Created {list_type} list with {len(items)} items"
    
    def _list_documents(self) -> str:
        """List all documents in the output directory."""
        try:
            files = list(self.output_dir.glob("*.docx"))
            if not files:
                return f"📁 No documents found in {self.output_dir}"
            
            file_list = []
            for file in files:
                stat = file.stat()
                size_kb = stat.st_size / 1024
                file_list.append(f"  📄 {file.name} ({size_kb:.1f} KB)")
            
            return f"📁 Documents in {self.output_dir}:\n" + "\n".join(file_list)
        
        except Exception as e:
            return f"❌ Failed to list documents: {str(e)}"

    async def handle_request(self, request: Dict[str, Any]) -> Dict[str, Any]:
        """Handle MCP requests."""
        method = request.get("method")
        request_id = request.get("id")
        params = request.get("params", {})
        
        try:
            if method == "initialize":
                return {
                    "jsonrpc": "2.0",
                    "id": request_id,
                    "result": {
                        "protocolVersion": "2024-11-05",
                        "capabilities": {
                            "tools": {}
                        },
                        "serverInfo": {
                            "name": "word-mcp-server-mac",
                            "version": "0.1.0"
                        }
                    }
                }
            
            elif method == "tools/list":
                return {
                    "jsonrpc": "2.0",
                    "id": request_id,
                    "result": {
                        "tools": self.tools
                    }
                }
            
            elif method == "tools/call":
                tool_name = params.get("name")
                arguments = params.get("arguments", {})
                
                # Execute the tool
                result_text = "❌ Unknown tool"
                
                if tool_name == "create_document":
                    title = arguments.get("title")
                    doc_id = self._create_document(title)
                    result_text = f"✅ Created document with ID: {doc_id}" + (f" and title: '{title}'" if title else "")
                
                elif tool_name == "insert_text":
                    doc_id = arguments.get("doc_id")
                    text = arguments.get("text")
                    result_text = self._insert_text(doc_id, text)
                
                elif tool_name == "format_text":
                    doc_id = arguments.get("doc_id")
                    bold = arguments.get("bold")
                    italic = arguments.get("italic")
                    font_size = arguments.get("font_size")
                    result_text = self._format_last_paragraph(doc_id, bold, italic, font_size)
                
                elif tool_name == "add_heading":
                    doc_id = arguments.get("doc_id")
                    text = arguments.get("text")
                    level = arguments.get("level", 1)
                    result_text = self._add_heading(doc_id, text, level)
                
                elif tool_name == "save_document":
                    doc_id = arguments.get("doc_id")
                    filename = arguments.get("filename")
                    result_text = self._save_document(doc_id, filename)
                
                elif tool_name == "read_document":
                    path = arguments.get("path")
                    result_text = self._read_document(path)
                
                elif tool_name == "create_table":
                    doc_id = arguments.get("doc_id")
                    rows = arguments.get("rows")
                    cols = arguments.get("cols")
                    result_text = self._create_table(doc_id, rows, cols)
                
                elif tool_name == "create_list":
                    doc_id = arguments.get("doc_id")
                    items = arguments.get("items")
                    list_type = arguments.get("list_type", "bulleted")
                    result_text = self._create_list(doc_id, items, list_type)
                
                elif tool_name == "list_documents":
                    result_text = self._list_documents()
                
                return {
                    "jsonrpc": "2.0",
                    "id": request_id,
                    "result": {
                        "content": [
                            {
                                "type": "text",
                                "text": result_text
                            }
                        ]
                    }
                }
            
            else:
                return {
                    "jsonrpc": "2.0",
                    "id": request_id,
                    "error": {
                        "code": -32601,
                        "message": f"Method not found: {method}"
                    }
                }
        
        except Exception as e:
            return {
                "jsonrpc": "2.0",
                "id": request_id,
                "error": {
                    "code": -32603,
                    "message": f"Internal error: {str(e)}"
                }
            }
    
    async def run(self):
        """Run the MCP server with stdio transport."""
        print(f"🍎 Mac Word MCP Server starting...", file=sys.stderr)
        print(f"📋 Registered {len(self.tools)} Word tools (Mac compatible)", file=sys.stderr)
        print(f"📁 Documents will be saved to: {self.output_dir}", file=sys.stderr)
        
        try:
            while True:
                # Read JSON-RPC request from stdin
                line = await asyncio.get_event_loop().run_in_executor(
                    None, sys.stdin.readline
                )
                
                if not line:
                    break
                
                try:
                    request = json.loads(line.strip())
                    response = await self.handle_request(request)
                    
                    # Send response to stdout
                    print(json.dumps(response), flush=True)
                    
                except json.JSONDecodeError as e:
                    print(f"❌ JSON decode error: {e}", file=sys.stderr)
                    continue
                    
        except KeyboardInterrupt:
            print("🛑 Server stopped by user", file=sys.stderr)
        except Exception as e:
            print(f"❌ Server error: {e}", file=sys.stderr)

async def main():
    """Main entry point."""
    server = MacWordMCPServer()
    await server.run()

if __name__ == "__main__":
    asyncio.run(main())